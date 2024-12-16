import tkinter as tk
from tkinter import ttk, messagebox
import sqlite3
from datetime import date, timedelta
from tkinter import simpledialog
from docx import Document
from datetime import datetime

# Veritabanı bağlantısı
def veritabani_baglanti():
    conn = sqlite3.connect("kafe_sistemi.db")
    cursor = conn.cursor()

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS kullanici (
        kullaniciID INTEGER PRIMARY KEY AUTOINCREMENT,
        kullaniciAdi TEXT NOT NULL,
        sifre TEXT NOT NULL,
        yetki TEXT NOT NULL
    )
    """)
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS kategori (
        kategoriID INTEGER PRIMARY KEY AUTOINCREMENT,
        kategoriAdi TEXT NOT NULL
    )
    """)
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS urun (
        urunID INTEGER PRIMARY KEY AUTOINCREMENT,
        urunAdi TEXT NOT NULL,
        kategoriID INTEGER,
        fiyat REAL NOT NULL,
        KDV REAL NOT NULL,
        stok INTEGER NOT NULL,
        FOREIGN KEY (kategoriID) REFERENCES kategori(kategoriID)
    )
    """)
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS masa (
        masano INTEGER PRIMARY KEY,
        masaDurumu TEXT NOT NULL
    )
    """)
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS siparis (
        siparisID INTEGER PRIMARY KEY AUTOINCREMENT,
        tarih TEXT NOT NULL,
        saat TEXT NOT NULL,
        masano INTEGER,
        garsonID INTEGER,
        toplamTutar REAL,
        KDV REAL,
        durum TEXT DEFAULT 'Bekliyor',
        FOREIGN KEY (masano) REFERENCES masa(masano),
        FOREIGN KEY (garsonID) REFERENCES kullanici(kullaniciID)
    )
    """)
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS siparis_alt (
        siparis_altID INTEGER PRIMARY KEY AUTOINCREMENT,
        siparisID INTEGER,
        urunID INTEGER,
        miktar INTEGER,
        fiyat REAL,
        KDV REAL,
        FOREIGN KEY (siparisID) REFERENCES siparis(siparisID),
        FOREIGN KEY (urunID) REFERENCES urun(urunID)
    )
    """)

     
    cursor.execute("SELECT * FROM kullanici WHERE kullaniciAdi = 'admin'")
    if not cursor.fetchone():
        cursor.execute("INSERT INTO kullanici (kullaniciAdi, sifre, yetki) VALUES ('admin', '123', 'Yönetici')")
        conn.commit()
    conn.close()


def yonetici_paneli():
    yonetici_pencere = tk.Tk()
    yonetici_pencere.title("Yönetici Paneli")
    yonetici_pencere.geometry("1600x700")

    # Ana çerçeve
    ana_frame = tk.Frame(yonetici_pencere, padx=20, pady=20)
    ana_frame.pack(fill="both", expand=True)

    # Garson Yönetimi
    sol_frame = tk.Frame(ana_frame, width=250, bg="lightgray")
    sol_frame.grid(row=0, column=0, padx=10, pady=10)

    tk.Label(sol_frame, text="Garson Yönetimi", bg="lightgray", font=("Arial", 14)).pack(pady=5)
    garson_listbox = tk.Listbox(sol_frame, width=25, height=10)
    garson_listbox.pack(padx=5, pady=5)

    tk.Label(sol_frame, text="Garson Adı:").pack(pady=5)
    entry_garson_adi = tk.Entry(sol_frame)
    entry_garson_adi.pack(pady=5)

    tk.Label(sol_frame, text="Şifre:").pack(pady=5)
    entry_garson_sifre = tk.Entry(sol_frame, show="*")
    entry_garson_sifre.pack(pady=5)

    tk.Button(sol_frame, text="Garson Ekle", command=lambda: garson_ekle(entry_garson_adi, entry_garson_sifre, garson_listbox)).pack(pady=5)
    tk.Button(sol_frame, text="Garson Sil", command=lambda: garson_sil(garson_listbox)).pack(pady=5)

    # Orta çerçeve: Ürün Yönetimi
    orta_frame = tk.Frame(ana_frame, width=250)
    orta_frame.grid(row=0, column=1, padx=10, pady=10)

    tk.Label(orta_frame, text="Ürün Yönetimi", font=("Arial", 14)).pack(pady=5)

    tk.Label(orta_frame, text="Ürün Adı:").pack(pady=5)
    entry_urun_adi = tk.Entry(orta_frame)
    entry_urun_adi.pack(pady=5)

    tk.Label(orta_frame, text="Kategori:").pack(pady=5)
    combo_kategori = ttk.Combobox(orta_frame)
    combo_kategori.pack(pady=5)

    tk.Label(orta_frame, text="Fiyat:").pack(pady=5)
    entry_fiyat = tk.Entry(orta_frame)
    entry_fiyat.pack(pady=5)

    tk.Label(orta_frame, text="KDV (%):").pack(pady=5)
    entry_kdv = tk.Entry(orta_frame)
    entry_kdv.pack(pady=5)

    tk.Label(orta_frame, text="Stok:").pack(pady=5)
    entry_stok = tk.Entry(orta_frame)
    entry_stok.pack(pady=5)

    tk.Button(orta_frame, text="Kategori Yükle", command=lambda: kategori_yukle(combo_kategori)).pack(pady=5)
    tk.Button(orta_frame, text="Ürün Ekle", command=lambda: urun_ekle(entry_urun_adi, combo_kategori, entry_fiyat, entry_kdv, entry_stok, urunler_listbox)).pack(pady=5)

    # Sağ çerçeve: Kategori Yönetimi
    sag_frame = tk.Frame(ana_frame, width=250, bg="lightgray")
    sag_frame.grid(row=0, column=2, padx=10, pady=10)

    tk.Label(sag_frame, text="Kategori Yönetimi", bg="lightgray", font=("Arial", 14)).pack(pady=5)
    kategori_listbox = tk.Listbox(sag_frame, width=25, height=10)
    kategori_listbox.pack(padx=5, pady=5)

    tk.Label(sag_frame, text="Kategori Adı:").pack(pady=5)
    entry_kategori_adi = tk.Entry(sag_frame)
    entry_kategori_adi.pack(pady=5)

    tk.Button(sag_frame, text="Kategori Ekle", command=lambda: kategori_ekle(entry_kategori_adi, kategori_listbox)).pack(pady=5)
    tk.Button(sag_frame, text="Kategori Sil", command=lambda: kategori_sil(kategori_listbox)).pack(pady=5)

    # Tedarikçi Yönetimi
    tedarikci_frame = tk.Frame(ana_frame, width=250, bg="lightblue")
    tedarikci_frame.grid(row=0, column=3, padx=10, pady=10)

    tk.Label(tedarikci_frame, text="Tedarikçi Yönetimi", bg="lightblue", font=("Arial", 14)).pack(pady=5)
    tedarikci_listbox = tk.Listbox(tedarikci_frame, width=30, height=10)
    tedarikci_listbox.pack(padx=5, pady=5)

    tk.Label(tedarikci_frame, text="Tedarikçi Adı:").pack(pady=5)
    entry_tedarikci_adi = tk.Entry(tedarikci_frame)
    entry_tedarikci_adi.pack(pady=5)

    tk.Label(tedarikci_frame, text="Telefon:").pack(pady=5)
    entry_tedarikci_telefon = tk.Entry(tedarikci_frame)
    entry_tedarikci_telefon.pack(pady=5)

    tk.Label(tedarikci_frame, text="Adres:").pack(pady=5)
    entry_tedarikci_adres = tk.Entry(tedarikci_frame)
    entry_tedarikci_adres.pack(pady=5)

    tk.Button(tedarikci_frame, text="Tedarikçi Ekle", command=lambda: tedarikci_ekle(entry_tedarikci_adi, entry_tedarikci_telefon, entry_tedarikci_adres, tedarikci_listbox)).pack(pady=5)
    tk.Button(tedarikci_frame, text="Tedarikçi Sil", command=lambda: tedarikci_sil(tedarikci_listbox)).pack(pady=5)

    # Raporlama Çerçevesi
    rapor_frame = tk.Frame(ana_frame, width=250, bg="white")
    rapor_frame.grid(row=0, column=4, padx=10, pady=10)

    tk.Label(rapor_frame, text="Raporlama", bg="white", font=("Arial", 14)).pack(pady=5)

    tk.Button(rapor_frame, text="Günlük Rapor", command=lambda: rapor_olustur("Günlük")).pack(pady=5)
    tk.Button(rapor_frame, text="Aylık Rapor", command=lambda: rapor_olustur("Aylık")).pack(pady=5)
    tk.Button(rapor_frame, text="Yıllık Rapor", command=lambda: rapor_olustur("Yıllık")).pack(pady=5)

    # Ürünler Çerçevesi 
    urunler_frame = tk.Frame(ana_frame, width=300, bg="lightyellow")
    urunler_frame.grid(row=0, column=5, padx=10, pady=10)

    tk.Label(urunler_frame, text="Mevcut Ürünler", bg="lightyellow", font=("Arial", 14)).pack(pady=5)

    urunler_listbox = tk.Listbox(urunler_frame, width=40, height=15)
    urunler_listbox.pack(padx=5, pady=5)

    tk.Label(urunler_frame, text="Ürün ID:").pack(pady=5)
    entry_urun_id = tk.Entry(urunler_frame)
    entry_urun_id.pack(pady=5)

    tk.Label(urunler_frame, text="Yeni Fiyat:").pack(pady=5)
    entry_yeni_fiyat = tk.Entry(urunler_frame)
    entry_yeni_fiyat.pack(pady=5)

    tk.Button(urunler_frame, text="Fiyat Güncelle", command=lambda: urun_fiyat_guncelle(entry_urun_id, entry_yeni_fiyat, urunler_listbox)).pack(pady=5)

    tk.Button(urunler_frame, text="Ürün Sil", command=lambda: urun_sil(entry_urun_id, urunler_listbox)).pack(pady=5)

    def rapor_olustur(tur):
        today = date.today()
        if tur == "Günlük":
            baslangic = today
            bitis = today
        elif tur == "Aylık":
            baslangic = today.replace(day=1)
            bitis = (today.replace(day=1) + timedelta(days=32)).replace(day=1) - timedelta(days=1)
        elif tur == "Yıllık":
            baslangic = today.replace(month=1, day=1)
            bitis = today.replace(month=12, day=31)

        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()

        query = """
        SELECT sa.siparis_altID, u.urunAdi, sa.miktar, sa.fiyat, sa.KDV 
        FROM siparis_alt sa
        JOIN urun u ON sa.urunID = u.urunID
        WHERE sa.siparis_altID IS NOT NULL
        """

        cursor.execute(query)
        results = cursor.fetchall()
        conn.close()

        document = Document()
        document.add_heading(f"{tur} Raporu", level=1)
        document.add_paragraph(f"Başlangıç Tarihi: {baslangic}")
        document.add_paragraph(f"Bitiş Tarihi: {bitis}")

        document.add_heading("Sipariş Detayları", level=2)
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sipariş Alt ID'
        hdr_cells[1].text = 'Ürün Adı'
        hdr_cells[2].text = 'Miktar'
        hdr_cells[3].text = 'Fiyat'
        hdr_cells[4].text = 'KDV'

        for row in results:
            row_cells = table.add_row().cells
            row_cells[0].text = str(row[0])
            row_cells[1].text = row[1]
            row_cells[2].text = str(row[2])
            row_cells[3].text = str(row[3])
            row_cells[4].text = str(row[4])

        file_name = f"{tur}_Raporu_{today}.docx"
        document.save(file_name)
        messagebox.showinfo("Rapor", f"{file_name} başarıyla oluşturuldu!")

    def urun_fiyat_guncelle(entry_id, entry_fiyat, listbox):
        urun_id = entry_id.get()
        fiyat = float(entry_fiyat.get())
        
        kdv_orani = 0.18

        kdv_tutari = fiyat * kdv_orani
        yeni_fiyat = fiyat + kdv_tutari

        if not urun_id or not yeni_fiyat:
            messagebox.showerror("Hata", "Tüm alanları doldurun!")
            return

        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()
        cursor.execute("UPDATE urun SET fiyat = ? WHERE urunID = ?", (yeni_fiyat, urun_id))
        conn.commit()
        conn.close()

        messagebox.showinfo("Başarılı", "Fiyat güncellendi!")
        urun_listele(listbox)

    def urun_sil(entry_id, listbox):
        urun_id = entry_id.get()

        if not urun_id:
            messagebox.showerror("Hata", "Ürün ID giriniz!")
            return

        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()
        cursor.execute("DELETE FROM urun WHERE urunID = ?", (urun_id,))
        conn.commit()
        conn.close()

        messagebox.showinfo("Başarılı", "Ürün silindi!")
        urun_listele(listbox)

    def urun_listele(listbox):
        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()
        cursor.execute("""
            SELECT u.urunID, u.urunAdi, u.fiyat, u.stok, k.kategoriAdi 
            FROM urun u
            LEFT JOIN kategori k ON u.kategoriID = k.kategoriID
        """)
        urunler = cursor.fetchall()
        conn.close()

        listbox.delete(0, tk.END)
        for urun in urunler:
            listbox.insert(tk.END, f"{urun[0]} - {urun[1]}: {urun[2]} TL, Stok: {urun[3]}, Kategori: {urun[4]}")

    def urun_ekle(entry_urun_adi, combo_kategori, entry_fiyat, entry_kdv, entry_stok, urunler_listbox):
        urun_adi = entry_urun_adi.get()
        kategori = combo_kategori.get()
        
        if not urun_adi or not kategori:
            messagebox.showerror("Hata", "Ürün adı ve kategori boş olamaz!")
            return

        try:
            fiyat = float(entry_fiyat.get())
            kdv = float(entry_kdv.get())
            stok = int(entry_stok.get())
            
            # KDV hesaplama
            kdv_tutari = fiyat * (kdv / 100)
            fiyat_kdv_dahil = fiyat + kdv_tutari
            
        except ValueError:
            messagebox.showerror("Hata", "Fiyat, KDV ve stok alanları geçerli bir değer içermelidir!")
            return

        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()
        
        # Kategori ID'sini al
        cursor.execute("SELECT kategoriID FROM kategori WHERE kategoriAdi = ?", (kategori,))
        kategori_id = cursor.fetchone()

        if not kategori_id:
            messagebox.showerror("Hata", "Lütfen önce kategori seçin!")
            conn.close()
            return

        try:
            cursor.execute("""
                INSERT INTO urun (urunAdi, kategoriID, fiyat, KDV, stok) 
                VALUES (?, ?, ?, ?, ?)
            """, (urun_adi, kategori_id[0], fiyat_kdv_dahil, kdv_tutari, stok))
            
            conn.commit()
            messagebox.showinfo("Başarılı", f"Ürün başarıyla eklendi!\nFiyat: {fiyat} TL\nKDV: {kdv_tutari:.2f} TL\nToplam: {fiyat_kdv_dahil:.2f} TL")
            
            # Alanları temizle
            entry_urun_adi.delete(0, tk.END)
            entry_fiyat.delete(0, tk.END)
            entry_kdv.delete(0, tk.END)
            entry_stok.delete(0, tk.END)
            combo_kategori.set('')
            
            # Ürün listesini güncelle
            urun_listele(urunler_listbox)
            
        except sqlite3.Error as e:
            messagebox.showerror("Hata", f"Veritabanı hatası: {e}")
        finally:
            conn.close()

    def tedarikci_ekle(entry_adi, entry_tel, entry_adres, listbox):
        ad = entry_adi.get()
        telefon = entry_tel.get()
        adres = entry_adres.get()

        if not ad or not telefon or not adres:
            messagebox.showerror("Hata", "Tüm alanları doldurun!")
            return

        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()
        cursor.execute("INSERT INTO tedarikci (unvan, telefon, adres) VALUES (?, ?, ?)", (ad, telefon, adres))
        conn.commit()
        conn.close()

        messagebox.showinfo("Başarılı", "Tedarikçi eklendi!")
        tedarikci_listele(listbox)

    def tedarikci_sil(listbox):
        secili_tedarikci = listbox.get(tk.ACTIVE)
        if not secili_tedarikci:
            messagebox.showerror("Hata", "Silmek için bir tedarikçi seçin!")
            return

        tedarikci_ad = secili_tedarikci.split(":")[0]

        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()
        cursor.execute("DELETE FROM tedarikci WHERE unvan = ?", (tedarikci_ad,))
        conn.commit()
        conn.close()

        messagebox.showinfo("Başarılı", "Tedarikçi silindi!")
        tedarikci_listele(listbox)

    def tedarikci_listele(listbox):
        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()
        cursor.execute("SELECT unvan, telefon FROM tedarikci")
        tedarikciler = cursor.fetchall()
        conn.close()

        listbox.delete(0, tk.END)
        for tedarikci in tedarikciler:
            listbox.insert(tk.END, f"{tedarikci[0]}: {tedarikci[1]}")
    tk.Button(yonetici_pencere, text="Çıkış", command=yonetici_pencere.destroy).pack(side="bottom", pady=10)
    
    # Masa Durumu Frame'i (Mevcut ürünlerin sağına eklenecek)
    masa_durum_frame = tk.Frame(ana_frame, width=250, bg="lightgreen")
    masa_durum_frame.grid(row=0, column=6, padx=10, pady=10)
    tk.Label(masa_durum_frame, text="Masa Durumları", bg="lightgreen", font=("Arial", 14)).pack(pady=5)

    # Masa durumlarını gösterecek treeview
    masa_tree = ttk.Treeview(masa_durum_frame, columns=("Masa No", "Durum"), show="headings", height=15)
    masa_tree.heading("Masa No", text="Masa No")
    masa_tree.heading("Durum", text="Durum")
    masa_tree.pack(padx=5, pady=5)

    def masa_durumlarini_goster():
        for item in masa_tree.get_children():
            masa_tree.delete(item)
            
        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()
        cursor.execute("SELECT masano, masaDurumu FROM masa")
        masalar = cursor.fetchall()
        conn.close()
        
        for masa in masalar:
            masa_tree.insert("", "end", values=masa)

    # Masa durumu güncelleme fonksiyonu
    def masa_durumu_guncelle():
        masa_no = masa_no_entry.get()
        yeni_durum = masa_durum_combo.get()
        
        if not masa_no or not yeni_durum:
            messagebox.showerror("Hata", "Lütfen tüm alanları doldurun!")
            return
            
        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()
        cursor.execute("UPDATE masa SET masaDurumu = ? WHERE masano = ?", (yeni_durum, masa_no))
        conn.commit()
        conn.close()
        
        messagebox.showinfo("Başarılı", "Masa durumu güncellendi!")
        masa_durumlarini_goster()

    # Yeni Masa Ekleme Penceresi
    def yeni_masa_ekle_pencere():
        masa_pencere = tk.Toplevel()
        masa_pencere.title("Yeni Masa Ekle")
        masa_pencere.geometry("300x150")
        masa_pencere.configure(bg="lightgreen")
        
        tk.Label(masa_pencere, text="Masa No:", bg="lightgreen").pack(pady=10)
        yeni_masa_no = tk.Entry(masa_pencere)
        yeni_masa_no.pack(pady=5)
        
        def masa_ekle():
            try:
                masa_no = int(yeni_masa_no.get())
                conn = sqlite3.connect("kafe_sistemi.db")
                cursor = conn.cursor()
                
                # Masa numarasının benzersiz olup olmadığını kontrol et
                cursor.execute("SELECT COUNT(*) FROM masa WHERE masano = ?", (masa_no,))
                if cursor.fetchone()[0] > 0:
                    messagebox.showerror("Hata", f"{masa_no} numaralı masa zaten mevcut!")
                    conn.close()
                    return
                    
                cursor.execute("INSERT INTO masa (masano, masaDurumu) VALUES (?, 'BOŞ')", (masa_no,))
                conn.commit()
                conn.close()
                
                messagebox.showinfo("Başarılı", f"{masa_no} numaralı masa eklendi!")
                masa_pencere.destroy()
                masa_durumlarini_goster()
            except ValueError:
                messagebox.showerror("Hata", "Geçerli bir masa numarası girin!")

        tk.Button(masa_pencere, text="Masa Ekle", command=masa_ekle).pack(pady=10)
        tk.Button(masa_pencere, text="İptal", command=masa_pencere.destroy).pack(pady=5)

    # Masa güncelleme kontrolleri
    tk.Label(masa_durum_frame, text="Masa No:", bg="lightgreen").pack(pady=5)
    masa_no_entry = tk.Entry(masa_durum_frame)
    masa_no_entry.pack(pady=5)

    tk.Label(masa_durum_frame, text="Yeni Durum:", bg="lightgreen").pack(pady=5)
    masa_durum_combo = ttk.Combobox(masa_durum_frame, values=["BOŞ", "DOLU", "REZERVE"])
    masa_durum_combo.pack(pady=5)

    tk.Button(masa_durum_frame, text="Güncelle", command=masa_durumu_guncelle).pack(pady=10)

    # Yeni Masa Ekle butonu
    tk.Button(masa_durum_frame, text="Yeni Masa Ekle", command=yeni_masa_ekle_pencere).pack(pady=10)

    masa_durumlarini_goster()

    kategori_yukle(combo_kategori)
    kategori_listele(kategori_listbox)
    garson_listele(garson_listbox)
    tedarikci_listele(tedarikci_listbox)
    urun_listele(urunler_listbox)

    yonetici_pencere.mainloop()


# Garson ekle fonksiyonu
def garson_ekle(entry_garson_adi, entry_garson_sifre, garson_listbox):
    garson_adi = entry_garson_adi.get()
    sifre = entry_garson_sifre.get()

    if not garson_adi or not sifre:
        messagebox.showerror("Hata", "Lütfen garson bilgilerini doldurun!")
        return

    conn = sqlite3.connect("kafe_sistemi.db")
    cursor = conn.cursor()
    cursor.execute("INSERT INTO kullanici (kullaniciAdi, sifre, yetki) VALUES (?, ?, 'Garson')", (garson_adi, sifre))
    conn.commit()
    conn.close()

    entry_garson_adi.delete(0, tk.END)
    entry_garson_sifre.delete(0, tk.END)
    messagebox.showinfo("Başarılı", "Garson başarıyla eklendi!")
    garson_listele(garson_listbox)

# Garson sil fonksiyonu
def garson_sil(garson_listbox):
    secili_garson = garson_listbox.get(tk.ACTIVE)
    if not secili_garson:
        messagebox.showerror("Hata", "Silmek için bir garson seçmelisiniz!")
        return

    conn = sqlite3.connect("kafe_sistemi.db")
    cursor = conn.cursor()
    cursor.execute("DELETE FROM kullanici WHERE kullaniciAdi = ? AND yetki = 'Garson'", (secili_garson,))
    conn.commit()
    conn.close()

    messagebox.showinfo("Başarılı", "Garson başarıyla silindi!")
    garson_listele(garson_listbox)

# Ürün ekle fonksiyonu
def urun_ekle(entry_urun_adi, combo_kategori, entry_fiyat, entry_kdv, entry_stok):
    urun_adi = entry_urun_adi.get()
    kategori = combo_kategori.get()

    try:
        fiyat = float(entry_fiyat.get())
        kdv = float(entry_kdv.get())
        stok = int(entry_stok.get())
    except ValueError:
        messagebox.showerror("Hata", "Fiyat, KDV ve stok alanları geçerli bir değer içermelidir!")
        return

    conn = sqlite3.connect("kafe_sistemi.db")
    cursor = conn.cursor()
    cursor.execute("SELECT kategoriID FROM kategori WHERE kategoriAdi = ?", (kategori,))
    kategori_id = cursor.fetchone()

    if not kategori_id:
        messagebox.showerror("Hata", "Kategori bulunamadı!")
        conn.close()
        return

    cursor.execute("INSERT INTO urun (urunAdi, kategoriID, fiyat, KDV, stok) VALUES (?, ?, ?, ?, ?)",
                   (urun_adi, kategori_id[0], fiyat, kdv, stok))
    conn.commit()
    conn.close()

    entry_urun_adi.delete(0, tk.END)
    entry_fiyat.delete(0, tk.END)
    entry_kdv.delete(0, tk.END)
    entry_stok.delete(0, tk.END)
    messagebox.showinfo("Başarılı", "Ürün başarıyla eklendi!")

# Kategori ekle fonksiyonu
def kategori_ekle(entry_kategori_adi, kategori_listbox):
    kategori_adi = entry_kategori_adi.get()

    if not kategori_adi:
        messagebox.showerror("Hata", "Kategori adı boş olamaz!")
        return

    conn = sqlite3.connect("kafe_sistemi.db")
    cursor = conn.cursor()
    cursor.execute("INSERT INTO kategori (kategoriAdi) VALUES (?)", (kategori_adi,))
    conn.commit()
    conn.close()

    entry_kategori_adi.delete(0, tk.END)
    messagebox.showinfo("Başarılı", "Kategori başarıyla eklendi!")
    kategori_listele(kategori_listbox)

# Kategori sil fonksiyonu
def kategori_sil(kategori_listbox):
    secili_kategori = kategori_listbox.get(tk.ACTIVE)

    if not secili_kategori:
        messagebox.showerror("Hata", "Silmek için bir kategori seçin!")
        return

    conn = sqlite3.connect("kafe_sistemi.db")
    cursor = conn.cursor()
    cursor.execute("DELETE FROM kategori WHERE kategoriAdi = ?", (secili_kategori,))
    conn.commit()
    conn.close()

    messagebox.showinfo("Başarılı", "Kategori başarıyla silindi!")
    kategori_listele(kategori_listbox)

# Garson listeleme
def garson_listele(garson_listbox):
    conn = sqlite3.connect("kafe_sistemi.db")
    cursor = conn.cursor()
    cursor.execute("SELECT kullaniciAdi FROM kullanici WHERE yetki = 'Garson'")
    garsonlar = [row[0] for row in cursor.fetchall()]
    conn.close()

    garson_listbox.delete(0, tk.END)
    for garson in garsonlar:
        garson_listbox.insert(tk.END, garson)

# Kategori listeleme
def kategori_listele(kategori_listbox):
    conn = sqlite3.connect("kafe_sistemi.db")
    cursor = conn.cursor()
    cursor.execute("SELECT kategoriAdi FROM kategori")
    kategoriler = [row[0] for row in cursor.fetchall()]
    conn.close()

    kategori_listbox.delete(0, tk.END)
    for kategori in kategoriler:
        kategori_listbox.insert(tk.END, kategori)

# Kategori yükle
def kategori_yukle(combo_kategori):
    conn = sqlite3.connect("kafe_sistemi.db")
    cursor = conn.cursor()
    cursor.execute("SELECT kategoriAdi FROM kategori")
    kategoriler = [row[0] for row in cursor.fetchall()]
    conn.close()

    combo_kategori['values'] = kategoriler


# Yardımcı fonksiyonlar
def kategori_yukle(combo_kategori):
    conn = sqlite3.connect("kafe_sistemi.db")
    cursor = conn.cursor()
    cursor.execute("SELECT kategoriAdi FROM kategori")
    kategoriler = [row[0] for row in cursor.fetchall()]
    conn.close()
    combo_kategori['values'] = kategoriler

def kategori_listele(kategori_listbox):
    conn = sqlite3.connect("kafe_sistemi.db")
    cursor = conn.cursor()
    cursor.execute("SELECT kategoriAdi FROM kategori")
    kategoriler = cursor.fetchall()
    conn.close()
    kategori_listbox.delete(0, tk.END)
    for kategori in kategoriler:
        kategori_listbox.insert(tk.END, kategori[0])

def garson_listele(garson_listbox):
    conn = sqlite3.connect("kafe_sistemi.db")
    cursor = conn.cursor()
    cursor.execute("SELECT kullaniciAdi FROM kullanici WHERE yetki = 'Garson'")
    garsonlar = cursor.fetchall()
    conn.close()
    garson_listbox.delete(0, tk.END)
    for garson in garsonlar:
        garson_listbox.insert(tk.END, garson[0])

def garson_paneli(garson_adi):
    garson_pencere = tk.Tk()
    garson_pencere.title(f"Garson Paneli - {garson_adi}")
    garson_pencere.geometry("1600x700")

    ana_frame = tk.Frame(garson_pencere, padx=20, pady=20)
    ana_frame.pack(fill="both", expand=True)
    tk.Button(garson_pencere, text="Çıkış", command=garson_pencere.destroy).pack(side="bottom", pady=10)

    # Sipariş Giriş Bölümü
    siparis_frame = tk.LabelFrame(ana_frame, text="Sipariş Gir", padx=10, pady=10)
    siparis_frame.grid(row=0, column=0, padx=10, pady=10, sticky="n")

    tk.Label(siparis_frame, text="Masa No:").grid(row=0, column=0, sticky="e")
    entry_masa_no = tk.Entry(siparis_frame)
    entry_masa_no.grid(row=0, column=1, padx=10)

    tk.Label(siparis_frame, text="Ürün:").grid(row=1, column=0, sticky="e")
    combo_urun = ttk.Combobox(siparis_frame)
    combo_urun.grid(row=1, column=1, padx=10)

    tk.Label(siparis_frame, text="Miktar:").grid(row=2, column=0, sticky="e")
    entry_miktar = tk.Entry(siparis_frame)
    entry_miktar.grid(row=2, column=1, padx=10)

    # Mevcut Ürünler Listesi
    urunler_frame = tk.LabelFrame(ana_frame, text="Mevcut Ürünler", padx=10, pady=10)
    urunler_frame.grid(row=0, column=1, rowspan=2, padx=10, pady=10, sticky="nsew")

    urunler_listbox = tk.Listbox(urunler_frame, width=30, height=10)
    urunler_listbox.pack(fill="both", expand=True)

    # Masa Durumu Frame'i
    masa_durum_frame = tk.LabelFrame(ana_frame, text="Masa Durumları", padx=10, pady=10)
    masa_durum_frame.grid(row=0, column=3, rowspan=3, padx=10, pady=10, sticky="nsew")

    columns_masa = ("Masa No", "Durum")
    masa_tablosu = ttk.Treeview(masa_durum_frame, columns=columns_masa, show="headings", height=20)
    
    for col in columns_masa:
        masa_tablosu.heading(col, text=col)
        masa_tablosu.column(col, width=100)
    
    masa_tablosu.pack(fill="both", expand=True)

    # Masa Birleştirme Bölümü
    birlestirme_frame = tk.LabelFrame(ana_frame, text="Masa Birleştirme", padx=10, pady=10)
    birlestirme_frame.grid(row=3, column=1, padx=10, pady=10, sticky="n")

    tk.Label(birlestirme_frame, text="Birleştirilecek Masalar:").grid(row=0, column=0, sticky="e")
    entry_masalar = tk.Entry(birlestirme_frame)
    entry_masalar.grid(row=0, column=1, padx=10)
    tk.Label(birlestirme_frame, text="(Örn: 1,2,3)").grid(row=1, column=1, sticky="w")

    def masa_birlestir():
        masalar = entry_masalar.get().split(',')
        if len(masalar) < 2:
            messagebox.showerror("Hata", "En az iki masa seçmelisiniz!")
            return
        
        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()
        
        hedef_masa = masalar[0]
        for masa in masalar[1:]:
            cursor.execute("UPDATE siparis SET masano = ? WHERE masano = ?", (hedef_masa, masa))
            cursor.execute("UPDATE masa SET masaDurumu = 'BOŞ' WHERE masano = ?", (masa,))
        
        cursor.execute("UPDATE masa SET masaDurumu = 'DOLU' WHERE masano = ?", (hedef_masa,))
        conn.commit()
        conn.close()
        
        messagebox.showinfo("Başarılı", f"Masalar {hedef_masa} numaralı masada birleştirildi!")
        siparisleri_goster()
        masa_durumlarini_goster()

    tk.Button(birlestirme_frame, text="Masaları Birleştir", command=masa_birlestir).grid(row=2, column=0, columnspan=2, pady=5)


    def masa_durumlarini_goster():
        for item in masa_tablosu.get_children():
            masa_tablosu.delete(item)
            
        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()
        cursor.execute("SELECT masano, masaDurumu FROM masa")
        masalar = cursor.fetchall()
        conn.close()
        
        for masa in masalar:
            masa_tablosu.insert("", "end", values=masa)

    def urun_yukle():
        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()
        cursor.execute("SELECT urunAdi,fiyat FROM urun")
        urunler = [(row[0], row[1]) for row in cursor.fetchall()]
        conn.close()
        combo_urun['values'] = urunler
        
        urunler_listbox.delete(0, tk.END)
        for urun in urunler:
            urunler_listbox.insert(tk.END, urun)

    def siparis_kaydet():
        masa_no = entry_masa_no.get()
        urun = combo_urun.get()
        miktar = entry_miktar.get()

        if not masa_no or not urun or not miktar:
            messagebox.showerror("Hata", "Lütfen tüm alanları doldurun!")
            return

        # Ürün adını fiyattan ayır
        urun_adi = urun.split()[0] if urun else ""

        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()

        # Önce ürünün var olup olmadığını kontrol et
        cursor.execute("SELECT EXISTS(SELECT 1 FROM urun WHERE urunAdi LIKE ?)", (f"{urun_adi}%",))
        urun_var = cursor.fetchone()[0]

        if not urun_var:
            messagebox.showerror("Hata", "Ürün bulunamadı!")
            conn.close()
            return

        # Ürün bilgilerini al
        cursor.execute("SELECT urunID, fiyat, KDV FROM urun WHERE urunAdi LIKE ?", (f"{urun_adi}%",))
        urun_bilgisi = cursor.fetchone()

        if urun_bilgisi:
            # Masa durumunu güncelle
            cursor.execute("UPDATE masa SET masaDurumu = 'DOLU' WHERE masano = ?", (masa_no,))

            # Sipariş oluştur
            cursor.execute("""
                INSERT INTO siparis (tarih, saat, masano, garsonID, durum) 
                VALUES (DATE('now'), TIME('now'), ?, 
                (SELECT kullaniciID FROM kullanici WHERE kullaniciAdi = ?), 'Bekliyor')
            """, (masa_no, garson_adi))
            siparis_id = cursor.lastrowid

            urun_id, fiyat, kdv = urun_bilgisi
            toplam_fiyat = fiyat * int(miktar)
            toplam_kdv = kdv * int(miktar)

            # Sipariş detaylarını ekle
            cursor.execute("""
                INSERT INTO siparis_alt (siparisID, urunID, miktar, fiyat, KDV) 
                VALUES (?, ?, ?, ?, ?)
            """, (siparis_id, urun_id, miktar, toplam_fiyat, toplam_kdv))

            # Sipariş toplamını güncelle
            cursor.execute("""
                UPDATE siparis SET toplamTutar = ?, KDV = ? 
                WHERE siparisID = ?
            """, (toplam_fiyat, toplam_kdv, siparis_id))

            conn.commit()
            messagebox.showinfo("Başarılı", "Sipariş başarıyla kaydedildi!")

            # Alanları temizle
            entry_masa_no.delete(0, tk.END)
            combo_urun.set('')
            entry_miktar.delete(0, tk.END)
        else:
            messagebox.showerror("Hata", "Ürün bilgileri alınamadı!")

        conn.close()
        siparisleri_goster()
        masa_durumlarini_goster()

    tk.Button(siparis_frame, text="Ürün Yükle", command=urun_yukle).grid(row=3, column=0, columnspan=2, pady=5)
    tk.Button(siparis_frame, text="Sipariş Kaydet", command=siparis_kaydet).grid(row=4, column=0, columnspan=2, pady=5)

    urun_yukle()

    # Sipariş Güncelle Bölümü
    guncelle_frame = tk.LabelFrame(ana_frame, text="Sipariş Güncelle", padx=10, pady=10)
    guncelle_frame.grid(row=1, column=0, padx=10, pady=10, sticky="n")

    tk.Label(guncelle_frame, text="Sipariş ID:").grid(row=0, column=0, sticky="e")
    entry_siparis_id = tk.Entry(guncelle_frame)
    entry_siparis_id.grid(row=0, column=1, padx=10)

    tk.Label(guncelle_frame, text="Yeni Durum:").grid(row=1, column=0, sticky="e")
    combo_durum = ttk.Combobox(guncelle_frame, values=["Bekliyor", "Hazırlanıyor", "Tamamlandı", "Ödendi"])
    combo_durum.grid(row=1, column=1, padx=10)

    def siparis_guncelle():
        siparis_id = entry_siparis_id.get()
        yeni_durum = combo_durum.get()

        if not siparis_id or not yeni_durum:
            messagebox.showerror("Hata", "Lütfen tüm alanları doldurun!")
            return

        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()
        cursor.execute("UPDATE siparis SET durum = ? WHERE siparisID = ?", (yeni_durum, siparis_id))
        conn.commit()
        conn.close()

        messagebox.showinfo("Başarılı", f"Sipariş {siparis_id} durumu '{yeni_durum}' olarak güncellendi!")
        siparisleri_goster()

    tk.Button(guncelle_frame, text="Güncelle", command=siparis_guncelle).grid(row=2, column=0, columnspan=2, pady=5)

    # Sipariş İptal Bölümü
    iptal_frame = tk.LabelFrame(ana_frame, text="Sipariş İptal", padx=10, pady=10)
    iptal_frame.grid(row=2, column=0, padx=10, pady=10, sticky="n")

    tk.Label(iptal_frame, text="Sipariş ID:").grid(row=0, column=0, sticky="e")
    entry_iptal_id = tk.Entry(iptal_frame)
    entry_iptal_id.grid(row=0, column=1, padx=10)

    def siparis_iptal():
        siparis_id = entry_iptal_id.get()

        if not siparis_id:
            messagebox.showerror("Hata", "Lütfen Sipariş ID'sini girin!")
            return

        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()
        
        # Siparişin masa numarasını al
        cursor.execute("SELECT masano FROM siparis WHERE siparisID = ?", (siparis_id,))
        masa_no = cursor.fetchone()
        
        if masa_no:
            # Masa durumunu BOŞ olarak güncelle
            cursor.execute("UPDATE masa SET masaDurumu = 'BOŞ' WHERE masano = ?", (masa_no[0],))
            
        cursor.execute("DELETE FROM siparis WHERE siparisID = ?", (siparis_id,))
        cursor.execute("DELETE FROM siparis_alt WHERE siparisID = ?", (siparis_id,))
        conn.commit()
        conn.close()

        messagebox.showinfo("Başarılı", f"Sipariş {siparis_id} iptal edildi!")
        siparisleri_goster()
        masa_durumlarini_goster()

    tk.Button(iptal_frame, text="Sipariş İptal Et", command=siparis_iptal).grid(row=1, column=0, columnspan=2, pady=5)

    # Sağ Çerçeve: Sipariş Tablosu
    siparis_tablosu_frame = tk.LabelFrame(ana_frame, text="Siparişler", padx=10, pady=10)
    siparis_tablosu_frame.grid(row=0, column=2, rowspan=3, padx=10, pady=10, sticky="nsew")

    columns = ("siparisID", "masaNo", "durum", "toplamTutar", "KDV")
    siparis_tablosu = ttk.Treeview(siparis_tablosu_frame, columns=columns, show="headings", height=20)

    for col in columns:
        siparis_tablosu.heading(col, text=col)
        siparis_tablosu.column(col, width=100)

    siparis_tablosu.pack(fill="both", expand=True)

    def siparisleri_goster():
        for i in siparis_tablosu.get_children():
            siparis_tablosu.delete(i)

        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()
        cursor.execute("SELECT siparisID, masano, durum, toplamTutar, KDV FROM siparis")
        siparisler = cursor.fetchall()
        conn.close()

        for siparis in siparisler:
            siparis_tablosu.insert("", "end", values=siparis)

    # Ödeme Alma Bölümü
    odeme_frame = tk.LabelFrame(ana_frame, text="Ödeme Al", padx=10, pady=10)
    odeme_frame.grid(row=3, column=0, padx=10, pady=10, sticky="n")

    tk.Label(odeme_frame, text="Sipariş ID:").grid(row=0, column=0, sticky="e")
    entry_odeme_id = tk.Entry(odeme_frame)
    entry_odeme_id.grid(row=0, column=1, padx=10)

    def odeme_tamamla():
        siparis_id = entry_odeme_id.get()

        if not siparis_id:
            messagebox.showerror("Hata", "Lütfen Sipariş ID'sini girin!")
            return

        conn = sqlite3.connect("kafe_sistemi.db")
        cursor = conn.cursor()

        # Siparişin varlığını kontrol et
        cursor.execute("SELECT masano FROM siparis WHERE siparisID = ?", (siparis_id,))
        masa_no = cursor.fetchone()

        if not masa_no:
            messagebox.showerror("Hata", "Geçersiz Sipariş ID!")
            conn.close()
            return

        odeme_pencere = tk.Toplevel()
        odeme_pencere.title("Sipariş Detayları")
        odeme_pencere.geometry("400x300")

        odeme_pencere.update_idletasks()
        screen_width = odeme_pencere.winfo_screenwidth()
        screen_height = odeme_pencere.winfo_screenheight()
        size = tuple(int(x) for x in odeme_pencere.geometry().split('+')[0].split('x'))
        x = (screen_width // 2) - (size[0] // 2)
        y = (screen_height // 2) - (size[1] // 2)
        odeme_pencere.geometry(f"{size[0]}x{size[1]}+{x}+{y}")

        # Ürünleri getir
        cursor.execute("SELECT urun.urunAdi, siparis_alt.miktar FROM siparis_alt JOIN urun ON siparis_alt.urunID = urun.urunID WHERE siparis_alt.siparisID = ?", (siparis_id,))
        urunler = cursor.fetchall()

        tk.Label(odeme_pencere, text="Sipariş Detayları", font=("Arial", 12, "bold")).pack(pady=5)

        for urun, miktar in urunler:
            tk.Label(odeme_pencere, text=f"{urun} - {miktar} adet").pack(anchor="w", padx=10)

        def odeme_onayla():
            cursor.execute("UPDATE masa SET masaDurumu = 'BOŞ' WHERE masano = ?", (masa_no[0],))
            cursor.execute("UPDATE siparis SET durum = 'Ödendi' WHERE siparisID = ?", (siparis_id,))
            conn.commit()
            messagebox.showinfo("Başarılı", f"Sipariş {siparis_id} ödemesi tamamlandı!")
            conn.close()
            odeme_pencere.destroy()
            siparisleri_goster()
            masa_durumlarini_goster()

        # Ödemeyi Tamamla butonu
        tk.Button(odeme_pencere, text="Ödemeyi Tamamla", command=odeme_onayla).pack(pady=10)

    tk.Button(odeme_frame, text="Ödeme Tamamla", command=odeme_tamamla).grid(row=1, column=0, columnspan=2, pady=5)

    # İlk yüklemede siparişleri ve masa durumlarını göster
    siparisleri_goster()
    masa_durumlarini_goster()

# Giriş ekranı
def giris_yap():
    kullanici_adi = entry_kullanici_adi.get()
    sifre = entry_sifre.get()
    conn = sqlite3.connect("kafe_sistemi.db")
    cursor = conn.cursor()
    cursor.execute("SELECT yetki FROM kullanici WHERE kullaniciAdi = ? AND sifre = ?", (kullanici_adi, sifre))
    result = cursor.fetchone()
    conn.close()
    if result:
        if result[0] == "Yönetici":
            root.destroy()
            yonetici_paneli()
        elif result[0] == "Garson":
            root.destroy()
            garson_paneli(kullanici_adi)
    else:
        messagebox.showerror("Hata", "Kullanıcı adı veya şifre yanlış!")

veritabani_baglanti()
root = tk.Tk()
root.title("Kafe Sistemi")
root.geometry("400x200")

window_width = 400
window_height = 200

screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

x_cordinate = int((screen_width / 2) - (window_width / 2))
y_cordinate = int((screen_height / 2) - (window_height / 2))
root.geometry(f"{window_width}x{window_height}+{x_cordinate}+{y_cordinate}")


tk.Label(root, text="Kullanıcı Adı:").grid(row=0, column=0, padx=10, pady=10)
entry_kullanici_adi = tk.Entry(root)
entry_kullanici_adi.grid(row=0, column=1, padx=10, pady=10)

tk.Label(root, text="Şifre:").grid(row=1, column=0, padx=10, pady=10)
entry_sifre = tk.Entry(root, show="*")
entry_sifre.grid(row=1, column=1, padx=10, pady=10)

tk.Button(root, text="Giriş Yap", command=giris_yap).grid(row=2, column=0, columnspan=2, pady=20)

root.mainloop()
